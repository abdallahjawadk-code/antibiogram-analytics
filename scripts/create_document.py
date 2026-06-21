# -*- coding: utf-8 -*-
import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Colors
TEAL = RGBColor(13, 148, 136)      # #0d9488
NAVY = RGBColor(15, 23, 42)       # #0f172a
DARK = RGBColor(30, 41, 59)       # #1e293b
MID = RGBColor(71, 85, 105)       # #475569
LIGHT_TEAL = RGBColor(240, 253, 250) # #f0fdfa
WHITE = RGBColor(255, 255, 255)

def set_rtl(paragraph):
    """Sets paragraph direction to Right-to-Left (RTL) for Arabic text."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    # Ensure text runs inside are also marked as RTL
    pPr_r = OxmlElement('w:rPr')
    r_rtl = OxmlElement('w:rtl')
    r_rtl.set(qn('w:val'), '1')
    pPr_r.append(r_rtl)

def add_heading_rtl(doc, text, level, space_before=12, space_after=6):
    """Adds an RTL heading with professional custom formatting."""
    p = doc.add_heading('', level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = TEAL
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = NAVY
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = MID
    return p

def add_paragraph_rtl(doc, text, font_size=11, bold=False, color=DARK, space_before=0, space_after=6, italic=False):
    """Adds a standard RTL paragraph with custom fonts and sizes."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return p

def add_bullet_rtl(doc, text, font_size=11, color=DARK):
    """Adds an RTL bullet point with custom styling."""
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(font_size)
    run.font.color.rgb = color
    return p

def add_callout_box(doc, title_ar, text_ar):
    """Creates a beautifully styled callout box with a teal left border."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Set cell width and padding
    cell = table.cell(0, 0)
    cell.width = Inches(6.0)
    
    # Shading and borders
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0FDFA') # Teal-50
    tcPr.append(shd)
    
    borders = OxmlElement('w:tcBorders')
    
    # Thick teal left border
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24') # 3pt
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), '0D9488')
    borders.append(left)
    
    # Remove others
    for border_name in ['top', 'bottom', 'right']:
        b = OxmlElement(f'w:{border_name}')
        b.set(qn('w:val'), 'none')
        borders.append(b)
        
    tcPr.append(borders)
    
    # Content
    p_title = cell.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p_title)
    p_title.paragraph_format.space_before = Pt(4)
    p_title.paragraph_format.space_after = Pt(2)
    run_title = p_title.add_run(title_ar)
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(11)
    run_title.font.bold = True
    run_title.font.color.rgb = TEAL
    
    p_body = cell.add_paragraph()
    p_body.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p_body)
    p_body.paragraph_format.space_before = Pt(0)
    p_body.paragraph_format.space_after = Pt(4)
    run_body = p_body.add_run(text_ar)
    run_body.font.name = 'Arial'
    run_body.font.size = Pt(10)
    run_body.font.color.rgb = DARK
    
    # Spacer
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(0)
    p_space.paragraph_format.space_after = Pt(6)

def create_document():
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # --- COVER PAGE ---
    # Top spacing
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(80)
    
    # App name
    p_brand = doc.add_paragraph()
    p_brand.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p_brand)
    run_brand = p_brand.add_run("عائلة برمجيات AntibioGram Pro")
    run_brand.font.name = 'Arial'
    run_brand.font.size = Pt(14)
    run_brand.font.bold = True
    run_brand.font.color.rgb = TEAL
    
    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p_title)
    p_title.paragraph_format.space_before = Pt(10)
    p_title.paragraph_format.space_after = Pt(5)
    run_title = p_title.add_run("الدليل التوضيحي الطبي والحسابي الشامل")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(28)
    run_title.font.bold = True
    run_title.font.color.rgb = NAVY
    
    # Document Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p_sub)
    p_sub.paragraph_format.space_after = Pt(150)
    run_sub = p_sub.add_run("شرح متكامل لآليات العمل السريرية، الحسابات الإحصائية، والمنطق البرمجي لنظام تحليلات الأنتيبايوجرام")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(13)
    run_sub.font.color.rgb = MID
    
    # Metadata
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p_meta)
    run_meta = p_meta.add_run("إعداد: فريق خبراء البرمجة الطبية والإحصاء الحيوي\nالتاريخ: يونيو 2026\nحالة المستند: معتمد سريرياً")
    run_meta.font.name = 'Arial'
    run_meta.font.size = Pt(10)
    run_meta.font.color.rgb = MID
    
    doc.add_page_break()
    
    # --- SECTION 1 ---
    add_heading_rtl(doc, "1. مقدمة عامة عن النظام", level=1)
    add_paragraph_rtl(doc, "يُعد نظام AntibioGram Pro أداة إكلينيكية برمجية متطورة تهدف إلى أتمتة وتحليل بيانات مزارع بكتيريا مقاومة مضادات الميكروبات (AMR) داخل المستشفيات والمرافق الصحية. يقدم النظام حلاً برمجياً متكاثفاً يتجاوز مجرد سرد الأرقام إلى استخلاص مؤشرات سريرية وإحصائية بالغة الدقة لتوجيه سياسات العلاج التجريبي وسياسات مكافحة العدوى.")
    
    add_paragraph_rtl(doc, "يهدف هذا المستند إلى توثيق كافة الآليات والقواعد الحسابية والبرمجية والسريرية المعتمدة في النظام ليكون بمثابة دليل مرجعي متكامل للأطباء، ومسؤولي مكافحة العدوى، ومطوري النظام الإحصائي الحيوي.")
    
    # --- SECTION 2 ---
    add_heading_rtl(doc, "2. الآليات والعمليات الرياضية والإحصائية", level=1)
    
    add_heading_rtl(doc, "أ. معادلة احتساب نسب الحساسية والمقاومة (S/I/R)", level=2)
    add_paragraph_rtl(doc, "تعتبر الدقة الحسابية هي الركيزة الأولى للأنتيبايوجرام التراكمي. يقوم البرنامج بحساب النسب ديناميكياً بناءً على إجمالي الفحوصات الفعالة لتجنب أي انحياز إحصائي:")
    add_bullet_rtl(doc, "حساب إجمالي الفحوصات المعتمدة (Denominator): Total = Susceptible + Intermediate + Resistant.")
    add_bullet_rtl(doc, "نسبة الحساسية (S%): قسمة عدد العزلات الحساسة على الإجمالي مضروباً في 100.")
    add_bullet_rtl(doc, "نسبة المقاومة الحقيقية (R%): تحسب حصراً بقسمة عزلات المقاومة الفعالة على الإجمالي، دون اللجوء للطرح التلقائي (100 - S%) لتفادي إدراج فئة Intermediate بشكل خاطئ كفئة مقاومة.")
    
    add_heading_rtl(doc, "ب. فاصل الثقة 95% بطريقة ويلسون (Wilson Score Interval)", level=2)
    add_paragraph_rtl(doc, "عندما يكون حجم العينات صغيراً في تقارير الأنتيبايوجرام، فإن فترات الثقة التقليدية (Wald Interval) تعطي نتائج مضللة أو غير حقيقية (مثل قيم سالبة أو أكبر من 100%). لذلك، يطبق البرنامج طريقة Wilson Score الإحصائية الموصى بها طبياً لضمان حدود ثقة دقيقة تقع دائماً ضمن النطاق الحقيقي [0%, 100%].")
    
    add_heading_rtl(doc, "ج. التغطية التجريبية الموزونة (WISCA)", level=2)
    add_paragraph_rtl(doc, "يتم استخدام خوارزمية ترجيحية مخصصة لتقدير التغطية التجريبية لمضاد حيوي معين ضد مجموعة من الميكروبات المسببة لمتلازمة مرضية معينة، حيث يتم إعطاء وزن نسبي (Weight) لكل بكتيريا بناءً على مدى انتشارها الفعلي في المستشفى (إجمالي عزلاتها)، مما يوفر للأطباء توجيهاً دقيقاً لوصف العلاج التجريبي الأكثر ملاءمة للمريض قبل ظهور نتائج مزرعته الخاصة.")
    
    add_heading_rtl(doc, "د. اختبارات المقارنة والتحليل الإحصائي (Chi-Square & Regression)", level=2)
    add_paragraph_rtl(doc, "يتضمن النظام عمليتين إحصائيتين أساسيتين للمقارنة والتنبؤ:")
    add_bullet_rtl(doc, "اختبار مربع كاي (Chi-Square Test): لمقارنة نسب المقاومة بين فترات زمنية أو منشآت صحية مختلفة، مع تطبيق تصحيح ييتس للاستمرارية (Yates' Correction) تلقائياً عند صغر حجم الخلايا المتوقعة لحماية النتائج من الدلالة الزائفة.")
    add_bullet_rtl(doc, "الانحدار الخطي البسيط (OLS): لتتبع اتجاهات المقاومة السنوية والتنبؤ الإحصائي بنسبة المقاومة للسنوات القادمة، مع اشتراط توفر بيانات لسنتين على الأقل لتفعيل التنبؤ.")
    
    # --- SECTION 3 ---
    add_heading_rtl(doc, "3. الآليات والعمليات الطبية والسريرية", level=1)
    
    add_heading_rtl(doc, "أ. معايير كسر التركيز العالمية (CLSI & EUCAST Breakpoints)", level=2)
    add_paragraph_rtl(doc, "يطبق البرنامج قواعد بيانات دقيقة للتركيز المثبط الأدنى (MIC Breakpoints) للتحقق سريرياً من تصنيف حساسية الميكروبات. يتيح النظام التبديل الفوري بين المعايير الأمريكية (CLSI) والأوروبية (EUCAST) مع إعادة التفسير التلقائي الفوري لتوزيع الـ MIC المخزن محلياً.")
    
    add_callout_box(doc, 
                    "مثال سريري هام:", 
                    "عزل بكتيريا E. coli واختبارها ضد مضاد Ceftriaxone بتركيز MIC = 2. تحت معيار CLSI يُصنف هذا التركيز كحساس متوسط (I)، بينما تحت معيار EUCAST يُصنف التركيز كـ مقاوم (R). البرنامج يقوم بالتحويل تلقائياً وإعادة حساب النسب في الأنتيبايوجرام بدقة بالغة.")
    
    add_heading_rtl(doc, "ب. عتبة الموثوقية الطبية (توصيات CLSI M39)", level=2)
    add_paragraph_rtl(doc, "يلتزم البرنامج بتوصية CLSI M39 الشهيرة والتي تمنع نشر أو اعتماد النسب المئوية التراكمية للحساسية لأي بكتيريا يقل إجمالي عينات الفحص لها عن 30 عزلة، حيث يقوم النظام بوضع إشارة تحذير واضحة تفيد بعدم موثوقية النسبة إحصائياً لحماية المرضى من العلاجات المبنية على عينات عشوائية غير كافية.")
    
    add_heading_rtl(doc, "ج. رصد الأنماط المظهرية للمقاومة (Superbugs Alerts)", level=2)
    add_paragraph_rtl(doc, "يحتوي البرنامج على خوارزمية ذكية لمراقبة مؤشرات انتشار الميكروبات الفائقة المقاومة وإطلاق تنبيهات الترصد الوبائي:")
    add_bullet_rtl(doc, "CRE (الأمعائيات المقاومة للكربابينيم).")
    add_bullet_rtl(doc, "MRSA (المكورات العنقودية المقاومة للميثيسيلين).")
    add_bullet_rtl(doc, "VRE / VRSA (المكورات المعوية والعنقودية المقاومة للفانكومايسين).")
    add_bullet_rtl(doc, "ESBL (مقاومة الجيل الثالث والرابع من السيفالوسبورينات).")
    
    # --- SECTION 4 ---
    add_heading_rtl(doc, "4. الآليات المنطقية والبرمجية", level=1)
    
    add_heading_rtl(doc, "أ. دمج البيانات التراكمي وتصفية التكرار", level=2)
    add_paragraph_rtl(doc, "لمنع الانحياز الإحصائي الناتج عن سحب مزارع متعددة لنفس المريض خلال فترة تنويمه، يقوم المنطق البرمجي بدمج العينات التكرارية وفق قاعدة العزلة الأولى لكل مريض (First Isolate Rule) مع تصفية المزارع التكرارية المتطابقة خلال الفترة التحليلية المحددة.")
    
    add_heading_rtl(doc, "ب. معالجة النصوص البرمجية للـ MIC", level=2)
    add_paragraph_rtl(doc, "يتم استخدام تعابير برمجية منتظمة (Regular Expressions) لتنقية قيم التركيز المدخلة نصياً وفك ترميز المعاملات الحسابية لضمان تحويلها لقيم عددية قابلة للمقارنة والفرز دون التسبب في أخطاء برمجية في بيئة التشغيل.")
    
    # --- SECTION 5 ---
    add_heading_rtl(doc, "5. الخلاصة والتوصيات", level=1)
    add_paragraph_rtl(doc, "يقدم نظام AntibioGram Pro نموذجاً يحتذى به في دمج البرمجة الحديثة مع الإحصاء الحيوي والعلوم الطبية المعقدة. إن دقته الرياضية العالية والتزامه الصارم بتوصيات CLSI M39 يجعلانه مرجعاً موثوقاً وآمناً لتوحيد وتطوير إحصاءات مقاومة المضادات الحيوية على مستوى المستشفيات أو الشبكات الإقليمية الموحدة.")
    
    # Save document
    desktop_path = os.path.expanduser("~/Desktop")
    file_path = os.path.join(desktop_path, "AntibioGram_Pro_Documentation.docx")
    
    # Fallback to OneDrive Desktop
    onedrive_desktop = os.path.join(os.path.expanduser("~"), "OneDrive", "Desktop")
    if os.path.exists(onedrive_desktop):
        file_path = os.path.join(onedrive_desktop, "AntibioGram_Pro_Documentation.docx")
        
    doc.save(file_path)
    print(f"Saved Word Doc to: {file_path}")

if __name__ == "__main__":
    create_document()
